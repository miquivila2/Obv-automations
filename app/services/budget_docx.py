"""Render a priced budget draft into a .docx in the Axo Capital format.

Reproduces the structure of Oblivion's gold-standard budget: a header, a plain-
language "short version" summary, and a month-by-month table of line items with a
subtotal. IVA and discounts are intentionally absent — a human adds those in the
CRM (locked decision, docs §5).

Kept deliberately thin and presentational: it takes an already-priced draft (the
math is done in budget_math) and lays it out. Returns the .docx as bytes, ready to
upload to Supabase Storage.
"""
from __future__ import annotations

from io import BytesIO
from itertools import groupby


def render_budget_docx(*, project_name: str, currency: str, lines: list[dict], subtotal: float) -> bytes:
    """Build the .docx and return its bytes. `lines` are priced line items from
    budget_math.price_line_items (each has category, description, hours, unit_rate,
    amount, month)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    doc.add_heading(f"{project_name} — Budget Proposal", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Currency: {currency} · Prepared by Oblivion · Confidential").italic = True

    # --- The short version -------------------------------------------------
    doc.add_heading("The short version", level=1)
    total_hours = sum(float(li["hours"]) for li in lines)
    doc.add_paragraph(
        f"This budget covers {total_hours:.0f} hours of work. Every line is hours × rate; "
        f"the subtotal below excludes IVA and any discounts, which are applied separately."
    )
    summary = doc.add_table(rows=0, cols=2)
    summary.style = "Light Grid Accent 1"
    for label, value in [("Total hours", f"{total_hours:.0f}"), (f"Subtotal ({currency})", f"{subtotal:,.2f}")]:
        cells = summary.add_row().cells
        cells[0].text = label
        cells[1].text = value

    # --- Month-by-month line items ----------------------------------------
    doc.add_heading("What you are paying for", level=1)

    def _month_key(li: dict) -> str:
        return li.get("month") or "Unscheduled"

    ordered = sorted(lines, key=lambda li: (li.get("position", 0)))
    for month, group in groupby(sorted(ordered, key=_month_key), key=_month_key):
        group = list(group)
        doc.add_heading(month, level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(["What we build", "Hours", f"Rate ({currency})", f"Cost ({currency})", "Category"]):
            table.rows[0].cells[i].text = head
        month_subtotal = 0.0
        for li in group:
            row = table.add_row().cells
            desc = li.get("description", "")
            just = li.get("justification", "")
            row[0].text = f"{desc}\n{just}".strip()
            row[1].text = f"{float(li['hours']):.0f}"
            row[2].text = f"{float(li['unit_rate']):,.2f}"
            row[3].text = f"{float(li['amount']):,.2f}"
            row[4].text = li.get("category", "")
            month_subtotal += float(li["amount"])
        sub = table.add_row().cells
        sub[0].text = "Subtotal"
        sub[3].text = f"{month_subtotal:,.2f}"

    # --- Total -------------------------------------------------------------
    doc.add_heading("Total", level=1)
    total_p = doc.add_paragraph()
    run = total_p.add_run(f"Subtotal (all periods, {currency}): {subtotal:,.2f}")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph("IVA and discounts, if any, are added in the CRM.").italic = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
