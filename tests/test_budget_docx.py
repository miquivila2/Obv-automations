"""render_budget_docx produces a real .docx — verified via python-docx that the
structure (header, month grouping, subtotal) reflects the priced draft. No
Supabase involved; this is pure rendering.
"""
from io import BytesIO

from docx import Document

from app.services.budget_docx import render_budget_docx

_LINES = [
    {
        "category": "Dev", "description": "Build X", "justification": "core feature",
        "hours": 10, "unit_rate": 20.0, "amount": 200.0, "month": "Month 1", "position": 0,
    },
    {
        "category": "QA", "description": "Test X", "justification": "coverage",
        "hours": 5, "unit_rate": 20.0, "amount": 100.0, "month": "Month 1", "position": 1,
    },
    {
        "category": "Dev", "description": "Build Y", "justification": "second feature",
        "hours": 8, "unit_rate": 20.0, "amount": 160.0, "month": "Month 2", "position": 2,
    },
]


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_returns_nonempty_valid_docx_bytes():
    data = render_budget_docx(project_name="Acme", currency="USD", lines=_LINES, subtotal=460.0)
    assert isinstance(data, bytes) and len(data) > 0
    Document(BytesIO(data))  # doesn't raise -> valid .docx


def test_header_has_project_name_and_currency():
    doc = Document(BytesIO(render_budget_docx(project_name="Acme", currency="USD", lines=_LINES, subtotal=460.0)))
    text = _all_text(doc)
    assert "Acme" in text
    assert "USD" in text


def test_groups_line_items_by_month():
    doc = Document(BytesIO(render_budget_docx(project_name="Acme", currency="USD", lines=_LINES, subtotal=460.0)))
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert "Month 1" in headings
    assert "Month 2" in headings


def test_subtotal_is_rendered():
    doc = Document(BytesIO(render_budget_docx(project_name="Acme", currency="USD", lines=_LINES, subtotal=460.0)))
    assert "460.00" in _all_text(doc)


def test_no_iva_or_discount_language_is_invented():
    # Locked decision (docs §5): IVA/discounts are added by a human in the CRM,
    # never computed here — the doc should say so, not silently omit it.
    doc = Document(BytesIO(render_budget_docx(project_name="Acme", currency="USD", lines=_LINES, subtotal=460.0)))
    assert "IVA" in _all_text(doc)
